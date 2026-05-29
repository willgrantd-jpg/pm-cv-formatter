"""
populate.py — PM CV Profile document builder (v2 design).

Opens PM_CV_Template_v2.docx as the base document so the header logo,
footer, page size and margins are preserved automatically.  Clears the
body and rebuilds it from structured data to exactly match the v2 design.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Brand constants ───────────────────────────────────────────────────────────
NAVY    = "1B2A4E"
GOLD    = "C8A961"
GRAY    = "6B7280"
BODY    = "1A1A1A"
CREAM   = "F4F1EA"
DIVIDER = "D8D5CC"
WHITE   = "FFFFFF"

COL_DATE    = 1800   # twips — date column
COL_CONTENT = 7560   # twips — content column
TBL_WIDTH   = 9360   # twips — full body width


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _set_spacing(para, before=0, after=0, line=None):
    """Set paragraph spacing in twips.  line=None leaves line-height at default."""
    ppr = para._p.get_or_add_pPr()
    for old in ppr.findall(qn('w:spacing')):
        ppr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'), str(after))
    if line is not None:
        sp.set(qn('w:line'), str(line))
        sp.set(qn('w:lineRule'), 'auto')
    ppr.append(sp)


def _set_para_border_bottom(para, color, sz, space):
    """Attach a bottom border to a paragraph."""
    ppr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:color'), color)
    bot.set(qn('w:sz'), str(sz))
    bot.set(qn('w:space'), str(space))
    pBdr.append(bot)
    ppr.append(pBdr)


def _add_run(para, text, font_name='Calibri', size_hp=18,
             bold=False, italic=False, color=BODY, char_spacing=None):
    """
    Add a run with fully explicit character formatting.
    size_hp  : font size in half-points (18 = 9 pt, 22 = 11 pt, etc.)
    char_spacing : tracking in 1/20 pt units (e.g. -10, 60, 150, 200)
    """
    run = para.add_run(text)
    rpr = run._r.get_or_add_rPr()

    # Font family
    rFonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:cs', 'w:eastAsia', 'w:hAnsi'):
        rFonts.set(qn(attr), font_name)
    rpr.append(rFonts)

    # Bold / italic
    if bold:
        rpr.append(OxmlElement('w:b'))
        rpr.append(OxmlElement('w:bCs'))
    if italic:
        rpr.append(OxmlElement('w:i'))
        rpr.append(OxmlElement('w:iCs'))

    # Colour
    col_el = OxmlElement('w:color')
    col_el.set(qn('w:val'), color)
    rpr.append(col_el)

    # Character tracking
    if char_spacing is not None:
        sp_el = OxmlElement('w:spacing')
        sp_el.set(qn('w:val'), str(char_spacing))
        rpr.append(sp_el)

    # Size
    for tag in ('w:sz', 'w:szCs'):
        el = OxmlElement(tag)
        el.set(qn('w:val'), str(size_hp))
        rpr.append(el)

    return run


def _set_table_width(table, width_twips):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(width_twips))
    tblPr.append(tblW)


def _no_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBdr = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:color'), WHITE)
        el.set(qn('w:sz'), '0')
        tblBdr.append(el)
    tblPr.append(tblBdr)


def _no_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:color'), WHITE)
        el.set(qn('w:sz'), '0')
        tcBdr.append(el)
    tcPr.append(tcBdr)


def _set_cell_width(cell, width_twips):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(width_twips))
    tcPr.append(tcW)


def _set_cell_margins(cell, top=0, left=0, bottom=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:type'), 'dxa')
        el.set(qn('w:w'), str(val))
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_cell_fill(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex6)
    tcPr.append(shd)


def _set_cell_border_left_only(cell, color, sz, space):
    """Gold left accent border; all other sides suppressed."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        if side == 'left':
            el.set(qn('w:val'), 'single')
            el.set(qn('w:color'), color)
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:space'), str(space))
        else:
            el.set(qn('w:val'), 'none')
            el.set(qn('w:color'), WHITE)
            el.set(qn('w:sz'), '0')
        tcBdr.append(el)
    tcPr.append(tcBdr)


def _set_tbl_grid(table, widths):
    """Replace tblGrid with explicit column widths (list of twips)."""
    tbl = table._tbl
    for old in tbl.findall(qn('w:tblGrid')):
        tbl.remove(old)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    # Insert after tblPr
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def _clear_body(doc):
    """Remove all body content; keep only sectPr so page setup is preserved."""
    body  = doc.element.body
    sect  = body.find(qn('w:sectPr'))
    for child in list(body):
        if child is not sect:
            body.remove(child)


def _fix_footer(doc, template_path):
    """
    Replace the small footer icon with the full PM logo PNG and fix margins.

    • Swaps the image bytes in the footer's image Part.
    • Sets display size to 1.2" wide (preserving 4957:1150 aspect ratio).
    • Sets paragraph line-spacing to atLeast so the image is never clipped.
    • Grows the bottom margin to 1800 twips so the footer zone is ~19 mm.
    """
    from pathlib import Path

    logo_path = Path(template_path).parent / "pm-logo.png"
    if not logo_path.exists():
        return

    # PM logo: 4957 × 1150 px — display at 1.2" wide
    LOGO_W = 1097280                              # 1.2" in EMU
    LOGO_H = int(LOGO_W * 1150 / 4957)           # ≈ 254 540 EMU ≈ 0.278"

    WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    A_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    W_NS  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # ── 1. Bottom margin ────────────────────────────────────────────────
    sectPr = doc.element.body.find(qn('w:sectPr'))
    if sectPr is not None:
        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is not None:
            pgMar.set(qn('w:bottom'), '1800')

    # ── 2. Footer image + extents ───────────────────────────────────────
    try:
        with open(logo_path, 'rb') as f:
            logo_bytes = f.read()

        for rel in doc.part.rels.values():
            if 'footer' not in rel.reltype:
                continue
            footer_part = rel.target_part
            ftr         = footer_part.element

            # Replace image blob in the footer's image relationship
            for img_rel in footer_part.rels.values():
                if 'image' in img_rel.reltype:
                    img_rel.target_part._blob = logo_bytes
                    break

            # Update wp:extent  and  a:ext (in spPr xfrm) to new dimensions
            for el in ftr.iter(f'{{{WP_NS}}}extent'):
                el.set('cx', str(LOGO_W))
                el.set('cy', str(LOGO_H))
            for el in ftr.iter(f'{{{A_NS}}}ext'):
                el.set('cx', str(LOGO_W))
                el.set('cy', str(LOGO_H))

            for p_el in ftr.findall(f'{{{W_NS}}}p'):
                has_drawing = any(True for _ in p_el.iter(f'{{{WP_NS}}}inline'))

                if has_drawing:
                    # Image paragraph — fix line-spacing so it is never clipped
                    pPr = p_el.find(f'{{{W_NS}}}pPr')
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        p_el.insert(0, pPr)
                    for old in pPr.findall(f'{{{W_NS}}}spacing'):
                        pPr.remove(old)
                    sp = OxmlElement('w:spacing')
                    sp.set(qn('w:before'), '0')
                    sp.set(qn('w:after'),  '60')
                    sp.set(qn('w:line'),   str(LOGO_H * 1440 // 914400 + 40))
                    sp.set(qn('w:lineRule'), 'atLeast')
                    pPr.append(sp)
                else:
                    # Text paragraph — remove the "Patrick Morgan  ·  " run,
                    # keep only the PAGE / NUMPAGES field runs
                    for r_el in list(p_el.findall(f'{{{W_NS}}}r')):
                        t_el = r_el.find(f'{{{W_NS}}}t')
                        if t_el is not None and 'Patrick Morgan' in (t_el.text or ''):
                            p_el.remove(r_el)

    except Exception:
        pass   # Never crash the document build over a footer issue


# ── Content builders ──────────────────────────────────────────────────────────

def _section_header(doc, title):
    """Bold navy section heading with bottom navy rule."""
    p = doc.add_paragraph()
    _set_spacing(p, before=360, after=120)
    _set_para_border_bottom(p, NAVY, sz=6, space=4)
    _add_run(p, title, font_name='Calibri', size_hp=20,
             bold=True, color=NAVY, char_spacing=60)


def _get_bullets(entry: dict) -> list:
    """Return bullets list from either new array format or legacy bullet1/2/3 fields."""
    if 'bullets' in entry and isinstance(entry.get('bullets'), list):
        return [b for b in entry['bullets'] if b]
    # Legacy fallback
    return [entry.get(f'bullet{i}') for i in range(1, 10)
            if entry.get(f'bullet{i}')]


def _selected_impact(doc, achievement):
    """Cream box with gold left border — Selected Impact panel."""
    if not achievement:
        return

    table = doc.add_table(rows=1, cols=1)
    _set_table_width(table, TBL_WIDTH)
    _no_table_borders(table)

    cell = table.rows[0].cells[0]
    _set_cell_width(cell, TBL_WIDTH)
    _set_cell_fill(cell, CREAM)
    _set_cell_border_left_only(cell, GOLD, sz=18, space=4)
    _set_cell_margins(cell, top=160, left=280, bottom=160, right=200)

    # "SELECTED IMPACT" label
    p_label = cell.paragraphs[0]
    _set_spacing(p_label, before=0, after=100)
    _add_run(p_label, 'KEY ACHIEVEMENTS', font_name='Montserrat',
             size_hp=16, bold=True, color=GOLD, char_spacing=150)

    title   = (achievement.get('title') or '').strip()
    bullets = _get_bullets(achievement)

    # Row 1: title (bold navy) + first bullet as gray context on same line
    if title:
        context = bullets[0] if bullets else ''
        p = cell.add_paragraph()
        _set_spacing(p, before=0, after=0 if len(bullets) <= 1 else 50, line=280)
        _add_run(p, title + ('  ' if context else ''),
                 font_name='Calibri', size_hp=20, color=NAVY)
        if context:
            _add_run(p, context, font_name='Calibri', size_hp=18, color=GRAY)

    # Remaining bullets as standalone bold navy lines
    for i, bullet in enumerate(bullets[1:], start=1):
        is_last = (i == len(bullets) - 1)
        p = cell.add_paragraph()
        _set_spacing(p, before=0, after=0 if is_last else 50, line=280)
        _add_run(p, bullet, font_name='Calibri', size_hp=20, color=NAVY)


def _two_col_table(doc):
    """Borderless 9360-wide table with 1800 / 7560 column split."""
    table = doc.add_table(rows=1, cols=2)
    _set_table_width(table, TBL_WIDTH)
    _no_table_borders(table)
    _set_tbl_grid(table, [COL_DATE, COL_CONTENT])
    # Remove the placeholder first row — rows are added per entry
    tr = table.rows[0]._tr
    tr.getparent().remove(tr)
    return table


def _add_experience_row(table, entry, is_last=False):
    """Append one experience entry row to the table."""
    row = table.add_row()

    # — Date cell —
    dc = row.cells[0]
    _no_cell_borders(dc)
    _set_cell_width(dc, COL_DATE)
    _set_cell_margins(dc, top=60, left=0, bottom=60, right=200)
    dp = dc.paragraphs[0]
    _set_spacing(dp, before=0, after=0)
    start = (entry.get('startYear') or '').strip()
    end   = (entry.get('endYear')   or '').strip()
    date_text = f"{start} — {end}" if (start or end) else ''
    _add_run(dp, date_text, size_hp=18, color=GRAY)

    # — Content cell —
    cc = row.cells[1]
    _no_cell_borders(cc)
    _set_cell_width(cc, COL_CONTENT)
    _set_cell_margins(cc, top=60, left=0, bottom=60, right=0)

    # Company name
    p_co = cc.paragraphs[0]
    _set_spacing(p_co, before=0, after=30)
    _add_run(p_co, entry.get('company') or '',
             font_name='Montserrat', size_hp=20, bold=True, color=NAVY)

    # Position (italic gold)
    p_pos = cc.add_paragraph()
    _set_spacing(p_pos, before=0, after=80)
    _add_run(p_pos, entry.get('position') or '',
             font_name='Montserrat', size_hp=18, italic=True, color=GOLD)

    # Bullets  —  gold em-dash prefix + body text (unlimited)
    for bullet in _get_bullets(entry):
        pb = cc.add_paragraph()
        _set_spacing(pb, before=0, after=40, line=280)
        _add_run(pb, '—  ', size_hp=18, color=GOLD)
        _add_run(pb, bullet, font_name='Montserrat', size_hp=18, color=BODY)

    # Thin warm-gray rule between entries (omit after the last one)
    if not is_last:
        ps = cc.add_paragraph()
        _set_spacing(ps, before=80, after=0)
        _set_para_border_bottom(ps, DIVIDER, sz=4, space=1)
        ps.add_run('')


def _add_experience_group_row(table, entry, is_last=False):
    """Append a multi-role (promoted) company entry — one company, multiple role blocks."""
    row = table.add_row()

    # — Date cell — overall span
    dc = row.cells[0]
    _no_cell_borders(dc)
    _set_cell_width(dc, COL_DATE)
    _set_cell_margins(dc, top=60, left=0, bottom=60, right=200)
    dp = dc.paragraphs[0]
    _set_spacing(dp, before=0, after=0)
    start = (entry.get('startYear') or '').strip()
    end   = (entry.get('endYear')   or '').strip()
    _add_run(dp, f"{start} — {end}" if (start or end) else '', size_hp=18, color=GRAY)

    # — Content cell —
    cc = row.cells[1]
    _no_cell_borders(cc)
    _set_cell_width(cc, COL_CONTENT)
    _set_cell_margins(cc, top=60, left=0, bottom=60, right=0)

    # Company name (once, at the top)
    p_co = cc.paragraphs[0]
    _set_spacing(p_co, before=0, after=40)
    _add_run(p_co, entry.get('company') or '', font_name='Montserrat', size_hp=20, bold=True, color=NAVY)

    roles = [r for r in (entry.get('roles') or []) if r]
    for r_idx, role in enumerate(roles):
        is_last_role = (r_idx == len(roles) - 1)

        # Position (italic gold) + date range (gray) on same line
        p_pos = cc.add_paragraph()
        _set_spacing(p_pos, before=20, after=60)
        _add_run(p_pos, role.get('position') or '', font_name='Montserrat', size_hp=18, italic=True, color=GOLD)
        r_start = (role.get('startYear') or '').strip()
        r_end   = (role.get('endYear')   or '').strip()
        if r_start or r_end:
            _add_run(p_pos, f'  {r_start} — {r_end}', size_hp=16, color=GRAY)

        # Bullets
        for bullet in _get_bullets(role):
            pb = cc.add_paragraph()
            _set_spacing(pb, before=0, after=40, line=280)
            _add_run(pb, '—  ', size_hp=18, color=GOLD)
            _add_run(pb, bullet, font_name='Montserrat', size_hp=18, color=BODY)

        # Thin divider between roles within the same company (not after last role)
        if not is_last_role:
            ps = cc.add_paragraph()
            _set_spacing(ps, before=60, after=0)
            _set_para_border_bottom(ps, DIVIDER, sz=4, space=1)
            ps.add_run('')

    # Thin warm-gray rule between companies (omit after the last entry)
    if not is_last:
        ps = cc.add_paragraph()
        _set_spacing(ps, before=80, after=0)
        _set_para_border_bottom(ps, DIVIDER, sz=4, space=1)
        ps.add_run('')


def _add_education_row(table, entry, is_last=False):
    """Append one education entry row to the table."""
    row = table.add_row()

    # — Date cell —
    dc = row.cells[0]
    _no_cell_borders(dc)
    _set_cell_width(dc, COL_DATE)
    _set_cell_margins(dc, top=60, left=0, bottom=60, right=200)
    dp = dc.paragraphs[0]
    _set_spacing(dp, before=0, after=0)
    start = (entry.get('startYear') or '').strip()
    end   = (entry.get('endYear')   or '').strip()
    date_text = f"{start} — {end}" if (start or end) else ''
    _add_run(dp, date_text, size_hp=18, color=GRAY)

    # — Content cell —
    cc = row.cells[1]
    _no_cell_borders(cc)
    _set_cell_width(cc, COL_CONTENT)
    _set_cell_margins(cc, top=60, left=0, bottom=60, right=0)

    # Institution
    p_inst = cc.paragraphs[0]
    _set_spacing(p_inst, before=0, after=30)
    _add_run(p_inst, entry.get('institution') or '',
             font_name='Montserrat', size_hp=20, bold=True, color=NAVY)

    # Degree (italic gold)
    p_deg = cc.add_paragraph()
    _set_spacing(p_deg, before=0, after=40)
    _add_run(p_deg, entry.get('degree') or '',
             font_name='Montserrat', size_hp=18, italic=True, color=GOLD)

    # Detail lines — honours, thesis, etc. (unlimited)
    for detail in _get_bullets(entry):
        pd = cc.add_paragraph()
        _set_spacing(pd, before=0, after=80, line=280)
        _add_run(pd, detail, size_hp=18, color=BODY)

    # Separator rule between entries
    if not is_last:
        ps = cc.add_paragraph()
        _set_spacing(ps, before=40, after=0)
        _set_para_border_bottom(ps, DIVIDER, sz=4, space=1)
        ps.add_run('')


def _add_certification_row(table, entry, is_last=False):
    """Append one certification entry row to the table."""
    row = table.add_row()

    # — Year cell —
    dc = row.cells[0]
    _no_cell_borders(dc)
    _set_cell_width(dc, COL_DATE)
    _set_cell_margins(dc, top=60, left=0, bottom=60, right=200)
    dp = dc.paragraphs[0]
    _set_spacing(dp, before=0, after=0)
    year = (entry.get('year') or '').strip()
    _add_run(dp, year, size_hp=18, color=GRAY)

    # — Content cell —
    cc = row.cells[1]
    _no_cell_borders(cc)
    _set_cell_width(cc, COL_CONTENT)
    _set_cell_margins(cc, top=60, left=0, bottom=60, right=0)

    # Certification name (bold navy, 11 pt)
    p_name = cc.paragraphs[0]
    _set_spacing(p_name, before=0, after=30)
    _add_run(p_name, entry.get('name') or '', font_name='Montserrat', size_hp=20, bold=True, color=NAVY)

    # Issuing body (italic gold, 9 pt) — only if present
    issuer = (entry.get('issuer') or '').strip()
    if issuer:
        p_iss = cc.add_paragraph()
        _set_spacing(p_iss, before=0, after=40)
        _add_run(p_iss, issuer, font_name='Montserrat', size_hp=18, italic=True, color=GOLD)

    # Thin separator between entries (omit after last)
    if not is_last:
        ps = cc.add_paragraph()
        _set_spacing(ps, before=40, after=0)
        _set_para_border_bottom(ps, DIVIDER, sz=4, space=1)
        ps.add_run('')


def _languages_skills(doc, languages, skills):
    """Languages row and Skills row in the standard two-column table."""
    table = _two_col_table(doc)

    # ── Languages ──
    lang_row = table.add_row()

    lbl = lang_row.cells[0]
    _no_cell_borders(lbl)
    _set_cell_width(lbl, COL_DATE)
    _set_cell_margins(lbl, top=60, left=0, bottom=60, right=200)
    _set_spacing(lbl.paragraphs[0], before=0, after=0)
    _add_run(lbl.paragraphs[0], 'Languages', size_hp=18, color=GRAY)

    lc = lang_row.cells[1]
    _no_cell_borders(lc)
    _set_cell_width(lc, COL_CONTENT)
    _set_cell_margins(lc, top=60, left=0, bottom=60, right=0)
    lp = lc.paragraphs[0]
    _set_spacing(lp, before=0, after=120, line=300)

    lang_list = [l for l in (languages or []) if l]
    for idx, lang in enumerate(lang_list):
        name    = (lang.get('name')    or '').strip()
        fluency = (lang.get('fluency') or '').strip()
        _add_run(lp, name + ' ', font_name='Montserrat', size_hp=18, bold=True, color=NAVY)
        _add_run(lp, f'({fluency})', font_name='Montserrat', size_hp=18, color=GRAY)
        if idx < len(lang_list) - 1:
            _add_run(lp, '  ·  ', size_hp=18, color=GRAY)

    # ── Skills ──
    sk_row = table.add_row()

    sk_lbl = sk_row.cells[0]
    _no_cell_borders(sk_lbl)
    _set_cell_width(sk_lbl, COL_DATE)
    _set_cell_margins(sk_lbl, top=60, left=0, bottom=60, right=200)
    _set_spacing(sk_lbl.paragraphs[0], before=0, after=0)
    _add_run(sk_lbl.paragraphs[0], 'Skills', size_hp=18, color=GRAY)

    sk_c = sk_row.cells[1]
    _no_cell_borders(sk_c)
    _set_cell_width(sk_c, COL_CONTENT)
    _set_cell_margins(sk_c, top=60, left=0, bottom=60, right=0)
    sp = sk_c.paragraphs[0]
    _set_spacing(sp, before=0, after=0, line=300)
    skills_text = '  ·  '.join(s for s in (skills or []) if s)
    _add_run(sp, skills_text, font_name='Montserrat', size_hp=18, color=BODY)


# ── LinkedIn badge ────────────────────────────────────────────────────────────

def _add_linkedin_badge(para, url: str = None):
    """
    Append a gold LinkedIn 'in' badge image to the name paragraph.
    Hyperlinked if url provided; plain placeholder if not — team can
    right-click → Edit Hyperlink in Word to add the URL manually.
    """
    from pathlib import Path
    from lxml import etree

    BADGE_PATH = Path(__file__).parent / "assets" / "linkedin-badge.png"
    if not BADGE_PATH.exists():
        return

    try:
        IMG_REL  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
        LINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

        with open(BADGE_PATH, 'rb') as f:
            img_bytes = f.read()

        # Add image to the document package
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        img_part = Part(PackURI('/word/media/linkedin_badge.png'), 'image/png', img_bytes)
        img_rid  = para.part.relate_to(img_part, IMG_REL)

        # Add URL hyperlink relationship if provided
        url_rid = None
        if url:
            url_rid = para.part.relate_to(url, LINK_REL, is_external=True)

        # Badge display size: 16 pt square
        SZ = 16 * 12700   # 203200 EMU

        # Namespace URIs
        WNS  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        WP   = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        ANS  = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        PIC  = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
        RELS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

        hlink_xml = (
            f'<a:hlinkClick xmlns:a="{ANS}" xmlns:r="{RELS}" r:id="{url_rid}"/>'
            if url_rid else ''
        )

        # Small space run before the badge
        sp_run = etree.fromstring(
            f'<w:r xmlns:w="{WNS}">'
            f'<w:rPr><w:sz w:val="22"/><w:szCs w:val="22"/></w:rPr>'
            f'<w:t xml:space="preserve">   </w:t>'
            f'</w:r>'
        )
        para._p.append(sp_run)

        # Inline image run
        run_xml = (
            f'<w:r xmlns:w="{WNS}">'
            f'<w:rPr><w:noProof/></w:rPr>'
            f'<w:drawing>'
            f'<wp:inline xmlns:wp="{WP}" distT="0" distB="0" distL="0" distR="0">'
            f'<wp:extent cx="{SZ}" cy="{SZ}"/>'
            f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
            f'<wp:docPr id="101" name="LinkedIn"/>'
            f'<wp:cNvGraphicFramePr>'
            f'<a:graphicFrameLocks xmlns:a="{ANS}" noChangeAspect="1"/>'
            f'</wp:cNvGraphicFramePr>'
            f'<a:graphic xmlns:a="{ANS}">'
            f'<a:graphicData uri="{PIC}">'
            f'<pic:pic xmlns:pic="{PIC}">'
            f'<pic:nvPicPr>'
            f'<pic:cNvPr id="101" name="LinkedIn">{hlink_xml}</pic:cNvPr>'
            f'<pic:cNvPicPr><a:picLocks noChangeAspect="1"/></pic:cNvPicPr>'
            f'</pic:nvPicPr>'
            f'<pic:blipFill>'
            f'<a:blip xmlns:r="{RELS}" r:embed="{img_rid}"/>'
            f'<a:stretch><a:fillRect/></a:stretch>'
            f'</pic:blipFill>'
            f'<pic:spPr>'
            f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{SZ}" cy="{SZ}"/></a:xfrm>'
            f'<a:prstGeom prst="roundRect"><a:avLst/></a:prstGeom>'
            f'<a:noFill/>'
            f'</pic:spPr>'
            f'</pic:pic>'
            f'</a:graphicData>'
            f'</a:graphic>'
            f'</wp:inline>'
            f'</w:drawing>'
            f'</w:r>'
        )
        para._p.append(etree.fromstring(run_xml.encode('utf-8')))

    except Exception:
        pass   # Never crash the document build over a badge


# ── Post-save: disable proofing in the output ZIP ────────────────────────────

def _suppress_proofing(docx_path: str):
    """
    Directly inject hideSpellingErrors + hideGrammaticalErrors into the saved
    docx ZIP — bypasses python-docx Settings API which doesn't always flush.
    """
    import zipfile, shutil, tempfile, re as _re

    inject = (
        b'<w:hideSpellingErrors w:val="1"/>'
        b'<w:hideGrammaticalErrors w:val="1"/>'
    )

    tmp = tempfile.mktemp(suffix='.docx')
    shutil.copy2(docx_path, tmp)

    with zipfile.ZipFile(tmp, 'r') as zin, \
         zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data_bytes = zin.read(item.filename)
            if item.filename == 'word/settings.xml':
                # Insert before closing </w:settings> tag
                data_bytes = data_bytes.replace(
                    b'</w:settings>', inject + b'</w:settings>'
                )
            zout.writestr(item, data_bytes)

    import os
    os.unlink(tmp)


# ── Structural notes parser ───────────────────────────────────────────────────

def _parse_structural(notes: str) -> dict:
    """
    Parse the structural notes textarea into a flags dict.

    Supported commands (case-insensitive):
        Photo: yes              → add passport-photo placeholder top-right
        Remove: achievements    → skip key achievements box
        Remove: certifications  → skip certifications section
        Remove: skills          → skip skills
        Remove: languages       → skip languages
    """
    flags = {
        'photo':                 False,
        'remove_achievements':   False,
        'remove_certifications': False,
        'remove_skills':         False,
        'remove_languages':      False,
    }
    if not notes or not notes.strip():
        return flags
    for line in notes.splitlines():
        s = line.strip().lower()
        if s.startswith('photo') and 'yes' in s:
            flags['photo'] = True
        elif s.startswith('remove'):
            val = s.split(':', 1)[-1].strip() if ':' in s else s[len('remove'):].strip()
            if 'achiev' in val:
                flags['remove_achievements'] = True
            elif 'certif' in val:
                flags['remove_certifications'] = True
            elif 'skill' in val:
                flags['remove_skills'] = True
            elif 'lang' in val:
                flags['remove_languages'] = True
    return flags


def _name_header(doc, data, show_photo: bool = False):
    """
    Renders candidate name row.

    show_photo=False (default):
        Name paragraph + LinkedIn badge, then full-width CANDIDATE PROFILE label.

    show_photo=True:
        Spacer → 2-col table (name left, photo right) → full-width CANDIDATE PROFILE
        label below the table so the divider spans the full page width.
    """
    if not show_photo:
        p_name = doc.add_paragraph()
        _set_spacing(p_name, before=0, after=60)
        _add_run(p_name, data.get('name') or 'Candidate Name',
                 font_name='Calibri', size_hp=56, bold=True,
                 color=NAVY, char_spacing=-10)
        _add_linkedin_badge(p_name, data.get('linkedin_url'))

        p_sub = doc.add_paragraph()
        _set_spacing(p_sub, before=0, after=200)
        _set_para_border_bottom(p_sub, NAVY, sz=8, space=6)
        _add_run(p_sub, 'CANDIDATE PROFILE',
                 font_name='Calibri', size_hp=18, color=GRAY, char_spacing=200)
        return

    # ── Photo variant ─────────────────────────────────────────────────────────
    # Layout:
    #   [spacer]
    #   ┌──────────────────────────┬──────────┐
    #   │  Name               [in] │          │
    #   │                          │  Photo   │
    #   └──────────────────────────┴──────────┘
    #   CANDIDATE PROFILE  ________________________  ← full-width divider

    PHOTO_W = 2000   # ~35 mm
    NAME_W  = TBL_WIDTH - PHOTO_W

    # Spacer so name doesn't sit hard against the header logo
    p_top = doc.add_paragraph()
    _set_spacing(p_top, before=0, after=160)
    p_top.add_run('')

    tbl = doc.add_table(rows=1, cols=2)

    # Borderless table, full body width
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(TBL_WIDTH))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Minimum row height ≈ 45 mm (passport photo proportions)
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), '2551')
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)

    row = tbl.rows[0]

    # ── Left cell: name + LinkedIn badge only ──
    left = row.cells[0]
    _no_cell_borders(left)
    _set_cell_width(left, NAME_W)
    _set_cell_margins(left, top=0, left=0, bottom=0, right=300)

    p_name = left.paragraphs[0]
    _set_spacing(p_name, before=0, after=0)
    _add_run(p_name, data.get('name') or 'Candidate Name',
             font_name='Calibri', size_hp=56, bold=True,
             color=NAVY, char_spacing=-10)
    _add_linkedin_badge(p_name, data.get('linkedin_url'))

    # ── Right cell: photo placeholder ──
    right = row.cells[1]
    _set_cell_width(right, PHOTO_W)
    _set_cell_margins(right, top=60, left=80, bottom=60, right=0)

    tcPr = right._tc.get_or_add_tcPr()

    # Light grey fill
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBEBEB')
    tcPr.append(shd)

    # Thin grey border
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(el)
    tcPr.append(tcBorders)

    # "Photo" text centred
    p_photo = right.paragraphs[0]
    _set_spacing(p_photo, before=900, after=0)
    pPr = p_photo._p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    _add_run(p_photo, 'Photo', size_hp=18, italic=True, color='AAAAAA')

    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

    # ── Full-width CANDIDATE PROFILE label + navy divider BELOW the table ──
    p_sub = doc.add_paragraph()
    _set_spacing(p_sub, before=80, after=200)
    _set_para_border_bottom(p_sub, NAVY, sz=8, space=6)
    _add_run(p_sub, 'CANDIDATE PROFILE',
             font_name='Calibri', size_hp=18, color=GRAY, char_spacing=200)


# ── Main entry point ──────────────────────────────────────────────────────────

def populate_template(template_path: str, data: dict, output_path: str,
                      structural_notes: str = ""):
    """
    Build a PM CV Profile .docx matching the v2 design.

    template_path : path to PM_CV_Template_v2.docx — preserves header/footer/images.
    data          : structured dict produced by extract.py.
    output_path   : destination .docx path.
    """
    flags = _parse_structural(structural_notes)

    doc = Document(template_path)
    _clear_body(doc)
    _fix_footer(doc, template_path)

    # 1 ── Candidate name (+ optional photo placeholder)
    _name_header(doc, data, show_photo=flags['photo'])

    # 2 ── Summary (Georgia italic, 10 pt)
    p_sum = doc.add_paragraph()
    _set_spacing(p_sum, before=0, after=240, line=340)
    _add_run(p_sum, data.get('summary') or '',
             font_name='Georgia', size_hp=20, italic=True, color=BODY)

    # 3 ── Key Achievements box
    if not flags['remove_achievements']:
        _selected_impact(doc, data.get('achievement'))

    # 4 ── Experience
    experience = [e for e in (data.get('experience') or []) if e]
    if experience:
        _section_header(doc, 'EXPERIENCE')
        exp_table = _two_col_table(doc)
        for i, entry in enumerate(experience):
            is_last = (i == len(experience) - 1)
            if entry.get('roles'):
                _add_experience_group_row(exp_table, entry, is_last=is_last)
            else:
                _add_experience_row(exp_table, entry, is_last=is_last)

    # 5 ── Education
    education = [e for e in (data.get('education') or []) if e]
    if education:
        _section_header(doc, 'EDUCATION')
        edu_table = _two_col_table(doc)
        for i, entry in enumerate(education):
            _add_education_row(edu_table, entry,
                               is_last=(i == len(education) - 1))

    # 6 ── Languages & Skills
    languages = [] if flags['remove_languages'] else (data.get('languages') or [])
    skills    = [] if flags['remove_skills']    else (data.get('skills')    or [])
    if languages or skills:
        _section_header(doc, 'LANGUAGES & SKILLS')
        _languages_skills(doc, languages, skills)

    # 7 ── Certifications
    certifications = [] if flags['remove_certifications'] else \
                     [c for c in (data.get('certifications') or []) if c]
    if certifications:
        _section_header(doc, 'CERTIFICATIONS')
        cert_table = _two_col_table(doc)
        for i, entry in enumerate(certifications):
            _add_certification_row(cert_table, entry,
                                   is_last=(i == len(certifications) - 1))

    doc.save(output_path)
    _suppress_proofing(output_path)
